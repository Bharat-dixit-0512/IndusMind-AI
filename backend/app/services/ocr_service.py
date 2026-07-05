import os
import logging
from typing import Optional
from PIL import Image

# Third-party libraries (imported inside functions or defensively)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text from PDF. Falls back to Tesseract OCR if text is sparse (scanned PDF).
    """
    text_content = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        
        for page_idx, page in enumerate(doc):
            page_text = page.get_text()
            # If text is very short/empty, run OCR on the page
            if len(page_text.strip()) < 50:
                logger.info(f"Page {page_idx} of {file_path} seems to be scanned. Running OCR...")
                try:
                    import pytesseract
                    # Render page as image
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    # Read image with PIL and OCR
                    from io import BytesIO
                    img = Image.open(BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(img)
                    text_content.append(ocr_text)
                except Exception as ocr_err:
                    logger.error(f"OCR failed for page {page_idx}: {ocr_err}")
                    text_content.append(page_text) # Fallback to whatever text was there
            else:
                text_content.append(page_text)
                
        doc.close()
    except ImportError:
        logger.warning("PyMuPDF (fitz) is not installed. Attempting fallback via pypdf...")
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                text_content.append(page.extract_text() or "")
        except Exception as e:
            logger.error(f"Failed to extract PDF using pypdf fallback: {e}")
            raise RuntimeError(f"No PDF parsing libraries available. {e}")
    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {e}")
        raise e

    return "\n\n--- Page Break ---\n\n".join(text_content)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts paragraphs from a Microsoft Word Document.
    """
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_cells))
        
        return "\n".join(paragraphs) + "\n\n" + "\n".join(table_text)
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}")
        raise e


def extract_text_from_xlsx(file_path: str) -> str:
    """
    Extracts tabular data from Excel, converting rows into structured text sentences
    so they retain context during vector search and RAG.
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheet_texts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_texts.append(f"Sheet: {sheet_name}")
            
            # Read rows
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
                
            # Assume first row has headers
            headers = [str(cell) if cell is not None else f"Column {i}" for i, cell in enumerate(rows[0])]
            
            for row_idx, row in enumerate(rows[1:], start=2):
                row_values = []
                for col_idx, cell in enumerate(row):
                    if cell is not None:
                        row_values.append(f"{headers[col_idx]}: {cell}")
                if row_values:
                    sheet_texts.append(f"Row {row_idx}: " + ", ".join(row_values))
                    
        return "\n".join(sheet_texts)
    except Exception as e:
        logger.error(f"Error reading XLSX {file_path}: {e}")
        raise e


def extract_text_from_image(file_path: str) -> str:
    """
    Extracts text from an image file using Tesseract OCR.
    """
    try:
        import pytesseract
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        logger.error(f"Error running OCR on image {file_path}: {e}")
        raise e


def extract_text(file_path: str) -> str:
    """
    Unified entry point to extract text from files based on their extension.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_text_from_xlsx(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
        return extract_text_from_image(file_path)
    elif ext in [".txt", ".md", ".json", ".csv"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
